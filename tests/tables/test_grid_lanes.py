"""XML-backed regression tests for omitted table grid coordinates."""

from io import BytesIO
import zipfile

from docx import Document as PythonDocxDocument
from lxml import etree

from docxnote import Cell, DocxDocument, Paragraph, Table


W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _document_with_grid_lanes(*, before=0, after=0, grid_width=5, rows=2):
    document = PythonDocxDocument()
    table = document.add_table(rows=rows, cols=3)
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            cell.text = f"R{row_index}C{cell_index}"

    stream = BytesIO()
    document.save(stream)
    files = {}
    with zipfile.ZipFile(BytesIO(stream.getvalue())) as source:
        for name in source.namelist():
            files[name] = source.read(name)

    xml = etree.fromstring(files["word/document.xml"])
    tbl = xml.find(f".//{W}tbl")
    assert tbl is not None
    tbl_grid = tbl.find(f"./{W}tblGrid")
    assert tbl_grid is not None
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for _ in range(grid_width):
        etree.SubElement(tbl_grid, f"{W}gridCol")

    for row in tbl.findall(f"./{W}tr"):
        tr_pr = etree.Element(f"{W}trPr")
        if before:
            etree.SubElement(tr_pr, f"{W}gridBefore").set(f"{W}val", str(before))
        if after:
            etree.SubElement(tr_pr, f"{W}gridAfter").set(f"{W}val", str(after))
        row.insert(0, tr_pr)

    files["word/document.xml"] = etree.tostring(
        xml, xml_declaration=True, encoding="UTF-8"
    )
    output = BytesIO()
    with zipfile.ZipFile(output, "w") as target:
        for name, data in files.items():
            target.writestr(name, data)
    return output.getvalue()


def _rewrite_table(data, row_specs, mutate=None):
    files = {}
    with zipfile.ZipFile(BytesIO(data)) as source:
        for name in source.namelist():
            files[name] = source.read(name)

    xml = etree.fromstring(files["word/document.xml"])
    tbl = xml.find(f".//{W}tbl")
    assert tbl is not None
    rows = tbl.findall(f"./{W}tr")
    for row, (before, after) in zip(rows, row_specs):
        tr_pr = row.find(f"./{W}trPr")
        if tr_pr is None:
            tr_pr = etree.Element(f"{W}trPr")
            row.insert(0, tr_pr)
        if before:
            etree.SubElement(tr_pr, f"{W}gridBefore").set(f"{W}val", str(before))
        if after:
            etree.SubElement(tr_pr, f"{W}gridAfter").set(f"{W}val", str(after))
    if mutate is not None:
        mutate(tbl)
    files["word/document.xml"] = etree.tostring(
        xml, xml_declaration=True, encoding="UTF-8"
    )
    output = BytesIO()
    with zipfile.ZipFile(output, "w") as target:
        for name, content in files.items():
            target.writestr(name, content)
    return output.getvalue()


def _cell_text(cell):
    return "\n".join(
        block.text for block in cell.blocks() if isinstance(block, Paragraph)
    )


class TestTableGridLanes:
    def test_standard_vertical_merge_preserves_following_physical_cells(self):
        document = PythonDocxDocument()
        table = document.add_table(rows=2, cols=3)
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                cell.text = f"R{row_index}C{cell_index}"
        table.cell(0, 0).merge(table.cell(1, 0)).text = "MERGED"

        stream = BytesIO()
        document.save(stream)
        parsed = DocxDocument.parse(stream.getvalue())
        table = next(block for block in parsed.blocks() if isinstance(block, Table))

        assert table.shape() == (2, 3)
        assert _cell_text(table[0, 0]) == "MERGED"
        assert table[0, 0] is table[1, 0]
        assert table[0, 0].bounds() == (0, 0, 2, 1)
        assert _cell_text(table[0, 1]) == "R0C1"
        assert _cell_text(table[0, 2]) == "R0C2"
        assert _cell_text(table[1, 1]) == "R1C1"
        assert _cell_text(table[1, 2]) == "R1C2"
        assert table[0, 1].bounds() == (0, 1, 1, 2)
        assert table[1, 1].bounds() == (1, 1, 2, 2)
        assert table[1, 2].bounds() == (1, 2, 2, 3)

    def test_grid_before_only_keeps_leading_synthetic_lane(self):
        document = DocxDocument.parse(_document_with_grid_lanes(before=1, grid_width=4))
        table = next(block for block in document.blocks() if isinstance(block, Table))

        assert table.shape() == (2, 4)
        assert _cell_text(table[0, 0]) == ""
        assert _cell_text(table[0, 1]) == "R0C0"
        assert table[0, 0].blocks() == ()
        assert table[0, 0].bounds() == (0, 0, 1, 1)
        assert table[0, 0].path == "t:0/r:0/c:0"
        assert document.resolve(table[0, 0].path).path == table[0, 0].path

    def test_grid_after_only_keeps_trailing_synthetic_lane(self):
        document = DocxDocument.parse(_document_with_grid_lanes(after=1, grid_width=4))
        table = next(block for block in document.blocks() if isinstance(block, Table))

        assert table.shape() == (2, 4)
        assert _cell_text(table[0, 2]) == "R0C2"
        assert _cell_text(table[0, 3]) == ""
        assert table[0, 3].bounds() == (0, 3, 1, 4)
        assert table[0, 3].path == "t:0/r:0/c:3"
        assert document.resolve(table[0, 3].path).path == table[0, 3].path

    def test_uneven_rows_keep_each_row_logical_coordinates(self):
        data = _document_with_grid_lanes(grid_width=5)

        def remove_last_cell(tbl):
            tbl.findall(f"./{W}tr")[0].remove(
                tbl.findall(f"./{W}tr")[0].findall(f"./{W}tc")[-1]
            )

        document = DocxDocument.parse(
            _rewrite_table(data, [(1, 2), (0, 2)], remove_last_cell)
        )
        table = next(block for block in document.blocks() if isinstance(block, Table))

        assert table.shape() == (2, 5)
        assert _cell_text(table[0, 0]) == ""
        assert _cell_text(table[0, 1]) == "R0C0"
        assert _cell_text(table[0, 2]) == "R0C1"
        assert _cell_text(table[0, 3]) == ""
        assert _cell_text(table[1, 2]) == "R1C2"
        assert _cell_text(table[1, 4]) == ""
        assert table[0, 3].bounds() == (0, 3, 1, 4)
        assert document.resolve(table[0, 3].path).path == table[0, 3].path

    def test_independent_grid_before_and_after_keep_synthetic_cells(self):
        data = _document_with_grid_lanes(before=1, after=1, grid_width=5)
        document = DocxDocument.parse(data)
        table = next(block for block in document.blocks() if isinstance(block, Table))

        assert table.shape() == (2, 5)
        assert _cell_text(table[0, 0]) == ""
        assert _cell_text(table[0, 1]) == "R0C0"
        assert _cell_text(table[1, 3]) == "R1C2"
        assert _cell_text(table[1, 4]) == ""
        assert table[0, 0].bounds() == (0, 0, 1, 1)
        assert table[0, 0].path == "t:0/r:0/c:0"
        assert document.resolve(table[0, 0].path).path == table[0, 0].path

    def test_spans_and_vertical_merges_follow_leading_offset(self):
        data = _document_with_grid_lanes(before=1, after=1, grid_width=6)
        xml = etree.fromstring(zipfile.ZipFile(BytesIO(data)).read("word/document.xml"))
        rows = xml.findall(f".//{W}tbl/{W}tr")

        first = rows[0].findall(f"./{W}tc")[0]
        tc_pr = first.find(f"./{W}tcPr")
        assert tc_pr is not None
        etree.SubElement(tc_pr, f"{W}gridSpan").set(f"{W}val", "2")
        etree.SubElement(tc_pr, f"{W}vMerge").set(f"{W}val", "restart")
        continuation = rows[1].findall(f"./{W}tc")[0]
        continuation_pr = continuation.find(f"./{W}tcPr")
        assert continuation_pr is not None
        etree.SubElement(continuation_pr, f"{W}gridSpan").set(f"{W}val", "2")
        etree.SubElement(continuation_pr, f"{W}vMerge")

        files = {}
        original = zipfile.ZipFile(BytesIO(data))
        for name in original.namelist():
            files[name] = original.read(name)
        files["word/document.xml"] = etree.tostring(
            xml, xml_declaration=True, encoding="UTF-8"
        )
        output = BytesIO()
        with zipfile.ZipFile(output, "w") as target:
            for name, content in files.items():
                target.writestr(name, content)

        document = DocxDocument.parse(output.getvalue())
        table = next(block for block in document.blocks() if isinstance(block, Table))
        merged = table[0, 1]

        assert table.shape() == (2, 6)
        assert table[0, 1] is table[0, 2] is table[1, 1] is table[1, 2]
        assert _cell_text(merged) == "R0C0"
        assert merged.bounds() == (0, 1, 2, 3)
        assert merged.path == "t:0/r:0/c:1"
        assert document.resolve(merged.path).path == merged.path
        assert isinstance(table[0, 0], Cell)
        assert _cell_text(table[0, 0]) == ""

    def test_vertical_merge_terminates_on_shifted_row_boundary(self):
        data = _document_with_grid_lanes(grid_width=4, rows=3)

        def add_merge(tbl):
            rows = tbl.findall(f"./{W}tr")
            first = rows[0].findall(f"./{W}tc")[0].find(f"./{W}tcPr")
            etree.SubElement(first, f"{W}vMerge").set(f"{W}val", "restart")
            continuation = rows[1].findall(f"./{W}tc")[1].find(f"./{W}tcPr")
            etree.SubElement(continuation, f"{W}vMerge")

        document = DocxDocument.parse(
            _rewrite_table(data, [(1, 0), (0, 1), (1, 0)], add_merge)
        )
        table = next(block for block in document.blocks() if isinstance(block, Table))

        merged = table[0, 1]
        assert table.shape() == (3, 4)
        assert table[0, 1] is table[1, 1]
        assert table[2, 1] is not merged
        assert merged.bounds() == (0, 1, 2, 2)
        assert table[2, 1].bounds() == (2, 1, 3, 2)
        assert merged.path == "t:0/r:0/c:1"
        assert document.resolve(merged.path).path == merged.path

    def test_vertical_merge_does_not_cross_omitted_lane(self):
        data = _document_with_grid_lanes(grid_width=4, rows=3)

        def add_merge(tbl):
            rows = tbl.findall(f"./{W}tr")
            first = rows[0].findall(f"./{W}tc")[0].find(f"./{W}tcPr")
            etree.SubElement(first, f"{W}vMerge").set(f"{W}val", "restart")
            continuation = rows[1].findall(f"./{W}tc")[0].find(f"./{W}tcPr")
            etree.SubElement(continuation, f"{W}vMerge")

        document = DocxDocument.parse(
            _rewrite_table(data, [(1, 0), (2, 0), (0, 1)], add_merge)
        )
        table = next(block for block in document.blocks() if isinstance(block, Table))

        merged = table[0, 1]
        assert table.shape() == (3, 4)
        assert merged.bounds() == (0, 1, 1, 2)
        assert table[1, 1].path == "t:0/r:1/c:1"
        assert table[1, 1] is not merged
        assert table[2, 1].path == "t:0/r:2/c:1"
