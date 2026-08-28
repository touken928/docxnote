"""测试批注阅读功能（基于文本视图）"""

from datetime import datetime, timezone
from io import BytesIO
import zipfile

from lxml import etree
from docx import Document as PythonDocxDocument

from docxnote import Comment, DocxDocument, Paragraph
from docxnote.namespaces import NS


def _first_paragraph_with_text(doc: DocxDocument) -> Paragraph:
    for block in doc.blocks():
        if isinstance(block, Paragraph) and (block.text or "").strip():
            return block
    raise AssertionError("no non-empty paragraph found")


class TestCommentReading:
    def test_stale_paragraph_wrapper_observes_new_comments(self, simple_doc):
        doc = DocxDocument.parse(simple_doc)
        first = _first_paragraph_with_text(doc)
        fresh = next(
            p for p in doc.blocks() if isinstance(p, Paragraph) and p.path == first.path
        )

        assert first.comments == ()
        fresh.comment("new", start=0, end=1, author="tester")

        assert [c.text for c in first.comments] == ["new"]

    def test_negative_and_oversized_range_across_hyperlink(self):
        pd_doc = PythonDocxDocument()
        paragraph = pd_doc.add_paragraph()
        paragraph.add_run("AB")
        part = paragraph.part
        rel_id = part.relate_to(
            "https://example.com",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = etree.Element(
            f"{{{NS['w']}}}hyperlink", {f"{{{NS['r']}}}id": rel_id}
        )
        run = etree.SubElement(hyperlink, f"{{{NS['w']}}}r")
        text = etree.SubElement(run, f"{{{NS['w']}}}t")
        text.text = "CDEF"
        paragraph._p.append(hyperlink)

        buffer = BytesIO()
        pd_doc.save(buffer)
        doc = DocxDocument.parse(buffer.getvalue())
        para = _first_paragraph_with_text(doc)

        created = para.comment("range", start=-4, end=999, author="tester")
        assert (created.start, created.end) == (2, 6)

        out = doc.render()
        doc2 = DocxDocument.parse(out, keep_comments=True)
        read_back = _first_paragraph_with_text(doc2).comments[0]
        assert (read_back.start, read_back.end) == (2, 6)

    def test_partial_comment_preserves_non_text_run_content(self):
        """局部批注拆分 run 时不应丢失非文本节点。"""
        pd_doc = PythonDocxDocument()
        paragraph = pd_doc.add_paragraph()
        run = paragraph.add_run("ABCD")
        run_el = run._r

        sym = etree.Element(f"{{{NS['w']}}}sym")
        sym.set(f"{{{NS['w']}}}font", "Wingdings")
        sym.set(f"{{{NS['w']}}}char", "F04A")
        run_el.append(sym)

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)

        doc = DocxDocument.parse(buffer.getvalue())
        para = _first_paragraph_with_text(doc)
        created = para.comment("inner", start=1, end=3, author="tester")
        assert (created.start, created.end) == (1, 3)

        out = doc.render()
        with zipfile.ZipFile(BytesIO(out)) as z:
            document_tree = etree.fromstring(z.read("word/document.xml"))
            symbols = document_tree.findall(".//w:sym", NS)
            assert len(symbols) == 1
            assert symbols[0].get(f"{{{NS['w']}}}font") == "Wingdings"
            assert symbols[0].get(f"{{{NS['w']}}}char") == "F04A"

        doc2 = DocxDocument.parse(out, keep_comments=True)
        para2 = _first_paragraph_with_text(doc2)
        matched = [
            c for c in para2.comments if c.text == "inner" and c.author == "tester"
        ]
        assert matched
        assert (matched[0].start, matched[0].end) == (1, 3)
        assert para2.text == "ABCD"

    def test_partial_comment_inside_hyperlink_roundtrips_exact_range(self):
        """超链接容器内的局部批注应能精确锚定并读回。"""
        pd_doc = PythonDocxDocument()
        paragraph = pd_doc.add_paragraph()
        part = paragraph.part
        rel_id = part.relate_to(
            "https://example.com",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = etree.Element(
            f"{{{NS['w']}}}hyperlink", {f"{{{NS['r']}}}id": rel_id}
        )
        run = etree.SubElement(hyperlink, f"{{{NS['w']}}}r")
        text = etree.SubElement(run, f"{{{NS['w']}}}t")
        text.text = "ClickHere"
        paragraph._p.append(hyperlink)

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)

        doc = DocxDocument.parse(buffer.getvalue())
        para = _first_paragraph_with_text(doc)
        assert para.text == "ClickHere"

        created = para.comment("link", start=0, end=5, author="tester")
        assert (created.start, created.end) == (0, 5)

        out = doc.render()
        doc2 = DocxDocument.parse(out, keep_comments=True)
        para2 = _first_paragraph_with_text(doc2)

        matched = [
            c for c in para2.comments if c.text == "link" and c.author == "tester"
        ]
        assert matched
        assert (matched[0].start, matched[0].end) == (0, 5)
        assert para2.text[matched[0].start : matched[0].end] == "Click"

    def test_partial_comment_inside_single_run_roundtrips_exact_range(self):
        """单个 run 中的局部批注应保留精确字符范围。"""
        pd_doc = PythonDocxDocument()
        pd_doc.add_paragraph("ABCDE")

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)

        doc = DocxDocument.parse(buffer.getvalue())
        para = _first_paragraph_with_text(doc)

        created = para.comment("inner", start=1, end=3, author="tester")
        assert (created.start, created.end) == (1, 3)

        out = doc.render()
        doc2 = DocxDocument.parse(out, keep_comments=True)
        para2 = _first_paragraph_with_text(doc2)

        matched = [
            c for c in para2.comments if c.text == "inner" and c.author == "tester"
        ]
        assert matched
        assert (matched[0].start, matched[0].end) == (1, 3)
        assert para2.text[matched[0].start : matched[0].end] == "BC"

    def test_multiple_partial_comments_on_one_paragraph_roundtrip_exact_ranges(self):
        """同一段话上的多次局部批注都应保留各自精确范围。"""
        pd_doc = PythonDocxDocument()
        pd_doc.add_paragraph("ABCDEFGHIJ")

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)

        doc = DocxDocument.parse(buffer.getvalue())
        para = _first_paragraph_with_text(doc)

        c1 = para.comment("first", start=1, end=3, author="u1")
        c2 = para.comment("second", start=4, end=7, author="u2")
        c3 = para.comment("third", start=7, end=9, author="u3")

        assert (c1.start, c1.end) == (1, 3)
        assert (c2.start, c2.end) == (4, 7)
        assert (c3.start, c3.end) == (7, 9)

        out = doc.render()
        doc2 = DocxDocument.parse(out, keep_comments=True)
        para2 = _first_paragraph_with_text(doc2)

        matched = {(c.text, c.author): (c.start, c.end) for c in para2.comments}
        assert matched[("first", "u1")] == (1, 3)
        assert matched[("second", "u2")] == (4, 7)
        assert matched[("third", "u3")] == (7, 9)
        assert para2.text[1:3] == "BC"
        assert para2.text[4:7] == "EFG"
        assert para2.text[7:9] == "HI"

    def test_comment_after_reparse_keeps_old_and_new_exact_ranges(self):
        """渲染后再次批注，同段落上的旧范围与新范围都应准确。"""
        pd_doc = PythonDocxDocument()
        pd_doc.add_paragraph("ABCDEFGHIJ")

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)

        doc1 = DocxDocument.parse(buffer.getvalue())
        para1 = _first_paragraph_with_text(doc1)
        created1 = para1.comment("first", start=1, end=3, author="u1")
        assert (created1.start, created1.end) == (1, 3)

        out1 = doc1.render()
        doc2 = DocxDocument.parse(out1, keep_comments=True)
        para2 = _first_paragraph_with_text(doc2)
        created2 = para2.comment("second", start=5, end=8, author="u2")
        assert (created2.start, created2.end) == (5, 8)

        out2 = doc2.render()
        doc3 = DocxDocument.parse(out2, keep_comments=True)
        para3 = _first_paragraph_with_text(doc3)

        matched = {(c.text, c.author): (c.start, c.end) for c in para3.comments}
        assert matched[("first", "u1")] == (1, 3)
        assert matched[("second", "u2")] == (5, 8)
        assert para3.text[1:3] == "BC"
        assert para3.text[5:8] == "FGH"

    def test_paragraph_comments_after_writing_and_reparse(self, simple_doc):
        """写入批注后重新解析，能够从段落上读出批注列表。"""
        doc = DocxDocument.parse(simple_doc)
        para = _first_paragraph_with_text(doc)
        original_text = para.text

        para.comment("测试批注", start=1, end=3, author="tester")

        out = doc.render()
        doc2 = DocxDocument.parse(out, keep_comments=True)
        para2 = _first_paragraph_with_text(doc2)

        comments = para2.comments
        assert isinstance(comments, tuple)
        assert len(comments) >= 1

        c0 = comments[0]
        assert isinstance(c0, Comment)
        assert c0.text == "测试批注"
        assert c0.author == "tester"
        assert c0.paragraph.text == original_text

        # 范围必须落在合法区间内
        assert 0 <= c0.start <= c0.end <= len(para2.text)

    def test_document_comments_collect_all(self, simple_doc):
        """DocxDocument.comments() 能返回文档中所有批注。"""
        doc = DocxDocument.parse(simple_doc)
        para = _first_paragraph_with_text(doc)

        para.comment("c1", start=0, end=2, author="u1")
        para.comment("c2", start=2, end=4, author="u2")

        out = doc.render()
        doc2 = DocxDocument.parse(out, keep_comments=True)

        comments = doc2.comments()
        assert isinstance(comments, tuple)
        # 至少包含我们刚刚添加的两个批注
        texts = {c.text for c in comments}
        assert "c1" in texts
        assert "c2" in texts

    def test_comments_respect_keep_comments_flag(self, simple_doc):
        """keep_comments=False 时不暴露原始批注；True 时可以读取。"""
        # 构造带批注的 DOCX
        doc = DocxDocument.parse(simple_doc)
        para = _first_paragraph_with_text(doc)
        para.comment("existing", author="orig")
        out = doc.render()

        # 默认：不保留批注
        doc_no = DocxDocument.parse(out)
        assert doc_no.comments() == ()
        p_no = _first_paragraph_with_text(doc_no)
        assert p_no.comments == ()

        # 显式保留：可以读取到
        doc_yes = DocxDocument.parse(out, keep_comments=True)
        comments_yes = doc_yes.comments()
        assert any(c.text == "existing" and c.author == "orig" for c in comments_yes)

    def test_comment_date_roundtrip_in_read_model(self, simple_doc):
        """自定义日期的批注，在读取模型中也能拿到同样的 datetime（UTC 时间）。"""
        fixed = datetime(2021, 1, 2, 3, 4, 5, tzinfo=timezone.utc)
        doc = DocxDocument.parse(simple_doc)
        para = _first_paragraph_with_text(doc)

        para.comment("dated", author="tester", date=fixed)
        out = doc.render()

        doc2 = DocxDocument.parse(out, keep_comments=True)
        para2 = _first_paragraph_with_text(doc2)
        cs = para2.comments
        assert len(cs) >= 1

        matched = [c for c in cs if c.text == "dated" and c.author == "tester"]
        assert matched
        # 使用 UTC 时区比较
        assert matched[0].date is not None
        assert matched[0].date.astimezone(timezone.utc) == fixed
