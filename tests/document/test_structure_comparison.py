"""测试文档结构与 python-docx 的一致性"""

import zipfile
from io import BytesIO

from docx import Document as PythonDocxDocument
from lxml import etree

from docxnote import DocxDocument, Paragraph, Table

W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _xml_body_sequence(doc_bytes: bytes) -> list[str]:
    """基于 word/document.xml 的 w:body 直接子元素生成真实 P/T 顺序。

    与 ``DocxDocument.blocks()`` 的语义对齐：只映射 ``w:p`` -> "P"、
    ``w:tbl`` -> "T"，忽略其他直接子元素（如 ``w:sectPr``）。
    """
    with zipfile.ZipFile(BytesIO(doc_bytes)) as zf:
        xml_bytes = zf.read("word/document.xml")
    root = etree.fromstring(xml_bytes)
    body = root.find("w:body", W_NS)
    assert body is not None, "word/document.xml 中未找到 w:body"
    sequence: list[str] = []
    for child in body:
        if not isinstance(child.tag, str):
            continue
        tag = etree.QName(child.tag).localname
        if tag == "p":
            sequence.append("P")
        elif tag == "tbl":
            sequence.append("T")
    return sequence


class TestStructureComparison:
    """测试文档结构解析的一致性"""

    def test_paragraph_order(self, simple_doc):
        """测试段落顺序与 python-docx 一致"""
        # docxnote
        dn_doc = DocxDocument.parse(simple_doc)
        dn_paragraphs = [b for b in dn_doc.blocks() if isinstance(b, Paragraph)]

        # python-docx
        pd_doc = PythonDocxDocument(BytesIO(simple_doc))
        pd_paragraphs = pd_doc.paragraphs

        # 段落数量应该一致
        assert len(dn_paragraphs) == len(pd_paragraphs)

        # 段落文本应该一致
        for dn_p, pd_p in zip(dn_paragraphs, pd_paragraphs):
            assert dn_p.text == pd_p.text

    def test_table_order(self, complex_doc):
        """测试表格顺序与 python-docx 一致"""
        # docxnote
        dn_doc = DocxDocument.parse(complex_doc)
        dn_blocks = list(dn_doc.blocks())

        # python-docx
        pd_doc = PythonDocxDocument(BytesIO(complex_doc))

        # 提取 docxnote 的表格
        dn_tables = [b for b in dn_blocks if isinstance(b, Table)]

        # python-docx 的表格
        pd_tables = pd_doc.tables

        # 表格数量应该一致
        assert len(dn_tables) == len(pd_tables)

    def test_mixed_blocks_order(self, complex_doc):
        """测试段落和表格混合顺序与 word/document.xml 的 w:body 一致"""
        # docxnote
        dn_doc = DocxDocument.parse(complex_doc)
        dn_blocks = list(dn_doc.blocks())

        # 记录 docxnote 的块类型序列
        dn_sequence: list[str] = []
        for block in dn_blocks:
            if isinstance(block, Paragraph):
                dn_sequence.append("P")
            elif isinstance(block, Table):
                dn_sequence.append("T")

        # 基于 word/document.xml 的 w:body 直接子元素生成真实 P/T 顺序
        xml_sequence = _xml_body_sequence(complex_doc)

        # 块数量必须一致
        assert len(dn_sequence) == len(
            xml_sequence
        ), f"块数量不一致: docxnote={len(dn_sequence)}, xml={len(xml_sequence)}"

        # 逐项比较顺序（不能只比较计数，否则"分组返回"等错误序列会漏检）
        for i, (dn_kind, xml_kind) in enumerate(zip(dn_sequence, xml_sequence)):
            assert (
                dn_kind == xml_kind
            ), f"第 {i} 个块顺序不一致: docxnote={dn_kind}, xml={xml_kind}"

    def test_empty_paragraphs_preserved(self, simple_doc):
        """测试空段落是否被保留"""
        # 创建包含空段落的文档
        pd_doc = PythonDocxDocument()
        pd_doc.add_paragraph("第一段")
        pd_doc.add_paragraph("")  # 空段落
        pd_doc.add_paragraph("第三段")

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # docxnote
        dn_doc = DocxDocument.parse(doc_bytes)
        dn_paragraphs = [b for b in dn_doc.blocks() if isinstance(b, Paragraph)]

        # python-docx
        pd_doc2 = PythonDocxDocument(BytesIO(doc_bytes))
        pd_paragraphs = pd_doc2.paragraphs

        # 数量应该一致
        assert len(dn_paragraphs) == len(pd_paragraphs)

        # 检查空段落
        assert dn_paragraphs[1].text == ""
        assert pd_paragraphs[1].text == ""
